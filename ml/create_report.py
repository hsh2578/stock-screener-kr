# -*- coding: utf-8 -*-
"""
박스권 돌파 ML 모델 보고서 생성 스크립트
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_report():
    doc = Document()

    # 기본 스타일 설정 (한글 폰트)
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    style.font.size = Pt(11)

    # 제목
    title = doc.add_heading('박스권 돌파 성공 예측 머신러닝 모델', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 부제목
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('주식 스크리너 AI 점수 산출 시스템')
    run.font.size = Pt(14)
    run.font.name = '맑은 고딕'

    doc.add_paragraph()

    # ========================================
    # 1. 프로젝트 개요
    # ========================================
    doc.add_heading('1. 프로젝트 개요', level=1)

    doc.add_heading('1.1 목적', level=2)
    p = doc.add_paragraph()
    p.add_run('박스권 돌파 패턴이 발생한 종목의 향후 수익률을 예측하여, 투자자에게 ')
    p.add_run('AI 기반 투자 판단 지표').bold = True
    p.add_run('를 제공하는 것을 목표로 합니다.')

    doc.add_heading('1.2 예측 목표', level=2)

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'

    headers = ['구분', '내용']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    data = [
        ['회귀 예측', '돌파 후 20거래일 내 최대 수익률 (%)'],
        ['분류 예측', '돌파 후 20거래일 내 15% 이상 상승 여부']
    ]
    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text

    doc.add_paragraph()

    # ========================================
    # 2. 데이터셋
    # ========================================
    doc.add_heading('2. 데이터셋', level=1)

    doc.add_heading('2.1 데이터 수집 조건', level=2)

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    headers = ['항목', '조건']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    data = [
        ['수집 기간', '2019년 4월 ~ 2025년 5월 (약 6년)'],
        ['시가총액 필터', '1,000억원 이상'],
        ['제외 종목', 'SPAC, ETF, ETN, 리츠, 인프라펀드'],
        ['박스권 기준', '40거래일 이상 횡보 (종가 변동폭 20% 이내)'],
        ['중복 돌파 간격', '최소 20거래일 (동일 종목 연속 돌파 제외)']
    ]
    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text

    doc.add_paragraph()

    doc.add_heading('2.2 데이터 통계', level=2)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    headers = ['항목', '값']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    data = [
        ['총 샘플 수', '14,567개'],
        ['성공 샘플 (15%↑)', '4,245개 (29.1%)'],
        ['실패 샘플', '10,322개 (70.9%)']
    ]
    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text

    doc.add_paragraph()

    # ========================================
    # 3. 피처 엔지니어링
    # ========================================
    doc.add_heading('3. 피처 엔지니어링', level=1)

    doc.add_heading('3.1 피처 구성 (총 15개)', level=2)

    # 피처 테이블
    table = doc.add_table(rows=16, cols=3)
    table.style = 'Table Grid'

    headers = ['카테고리', '피처명', '설명']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    features = [
        ['돌파 품질', 'breakout_strength', '(종가 - 저항선) / 저항선'],
        ['', 'breakout_gap', '돌파일 시가갭 비율'],
        ['', 'close_strength', '(종가 - 저가) / (고가 - 저가)'],
        ['거래량', 'volume_surge', '당일거래량 / 20일 평균거래량'],
        ['', 'volume_dry_up', '최근10일 평균 / 50일 평균 (거래량 감소)'],
        ['', 'liquidity', '5일 평균 거래대금'],
        ['박스권', 'box_range_pct', '박스권 내 변동폭 (%)'],
        ['추세', 'ma20_deviation', '종가 / 20일선 - 1'],
        ['', 'ma200_slope', '200일선 기울기 (20일 변화율)'],
        ['', 'pct_above_52w_low', '52주 저가 대비 상승률'],
        ['변동성', 'volatility_contraction', '최근5일 변동폭 / 20일 변동폭'],
        ['', 'atr_ratio', 'ATR(14) / 종가'],
        ['상대강도', 'rs_vs_market', '종목 20일 수익률 - 시장 수익률'],
        ['', 'market_return', 'KOSPI/KOSDAQ 20일 수익률'],
        ['시간', 'days_since_ath', '역대 신고가 이후 경과일']
    ]

    for i, row_data in enumerate(features):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text

    doc.add_paragraph()

    doc.add_heading('3.2 피처셋 구성', level=2)

    p = doc.add_paragraph()
    p.add_run('회귀 모델 (10개 피처)').bold = True
    p.add_run(': ma20_deviation, pct_above_52w_low, breakout_strength, close_strength, ')
    p.add_run('days_since_ath, atr_ratio, liquidity, breakout_gap, rs_vs_market, market_return')

    p = doc.add_paragraph()
    p.add_run('분류 모델 (15개 피처)').bold = True
    p.add_run(': 위 10개 + volume_surge, box_range_pct, volume_dry_up, ma200_slope, volatility_contraction')

    doc.add_paragraph()

    # ========================================
    # 4. 모델 선정
    # ========================================
    doc.add_heading('4. 모델 선정 및 비교', level=1)

    doc.add_heading('4.1 회귀 모델 비교 (R² 기준)', level=2)

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'

    headers = ['모델', 'R² Score', 'MAE']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    reg_models = [
        ['Ridge', '0.0675', '11.27'],
        ['Lasso', '0.0677', '11.27'],
        ['Random Forest', '0.0408', '11.35'],
        ['XGBoost', '0.0370', '11.39'],
        ['LightGBM', '0.0569', '11.34'],
        ['Voting Ensemble', '0.0587', '11.27'],
        ['Stacking Ensemble', '0.0658', '11.25']
    ]

    for i, row_data in enumerate(reg_models):
        for j, text in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = text
            if i == 0:  # Ridge 강조
                set_cell_shading(cell, 'E2EFDA')

    p = doc.add_paragraph()
    p.add_run('→ 최종 선택: Ridge (alpha=10.0)').bold = True
    p.add_run(' - 가장 안정적인 성능, 과적합 방지')

    doc.add_paragraph()

    doc.add_heading('4.2 분류 모델 비교 (F1 Score 기준)', level=2)

    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'

    headers = ['모델', 'F1 Score', 'Precision', 'AUC']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    cls_models = [
        ['Logistic (Ridge)', '0.216', '0.572', '0.661'],
        ['Logistic (Lasso)', '0.216', '0.573', '0.661'],
        ['Random Forest', '0.192', '0.613', '0.680'],
        ['XGBoost', '0.282', '0.562', '0.680'],
        ['LightGBM', '0.271', '0.546', '0.680'],
        ['Voting Ensemble', '0.240', '0.590', '0.685'],
        ['Stacking Ensemble', '0.190', '0.631', '0.684']
    ]

    for i, row_data in enumerate(cls_models):
        for j, text in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = text
            if i == 3:  # XGBoost 강조
                set_cell_shading(cell, 'E2EFDA')

    p = doc.add_paragraph()
    p.add_run('→ 최종 선택: XGBoost').bold = True
    p.add_run(' - 가장 높은 F1 Score, 불균형 데이터 처리 우수')

    doc.add_paragraph()

    # ========================================
    # 5. 최종 모델 성능
    # ========================================
    doc.add_heading('5. 최종 모델 성능', level=1)

    doc.add_heading('5.1 회귀 모델 (Ridge)', level=2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ['지표', '값']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    data = [
        ['R² Score', '0.0687 (6.87%)'],
        ['MAE (평균 절대 오차)', '11.27%'],
        ['RMSE (평균 제곱근 오차)', '20.40%'],
        ['하이퍼파라미터', 'alpha = 10.0']
    ]
    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text

    doc.add_paragraph()

    doc.add_heading('5.2 분류 모델 (XGBoost)', level=2)

    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'

    headers = ['지표', '값']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    data = [
        ['Accuracy', '0.7103 (71.03%)'],
        ['Precision', '0.5062 (50.62%)'],
        ['Recall', '0.2311 (23.11%)'],
        ['F1 Score', '0.3174 (31.74%)'],
        ['AUC', '0.6670 (66.70%)'],
        ['하이퍼파라미터', 'n_estimators=200, max_depth=7, learning_rate=0.1']
    ]
    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('성능 해석:').bold = True

    doc.add_paragraph('• Precision 50.62%: 모델이 "성공"으로 예측한 종목 중 절반이 실제로 15% 이상 상승', style='List Bullet')
    doc.add_paragraph('• Recall 23.11%: 실제 성공 종목 중 약 23%를 포착 (보수적 예측)', style='List Bullet')
    doc.add_paragraph('• AUC 0.667: 랜덤(0.5) 대비 약 33% 향상된 분류 능력', style='List Bullet')

    doc.add_paragraph()

    # ========================================
    # 6. AI 점수 산출
    # ========================================
    doc.add_heading('6. AI 점수 산출 방식', level=1)

    doc.add_heading('6.1 점수 계산 공식', level=2)

    p = doc.add_paragraph()
    run = p.add_run('AI Score = 성공확률 × 0.7 + 수익점수 × 0.3')
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    headers = ['구성요소', '설명']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    data = [
        ['성공확률', 'XGBoost 분류 모델의 예측 확률 (0~1)'],
        ['수익점수', '예상수익률을 0~1로 정규화 (상한 50%, 하한 -20%)'],
        ['가중치', '성공확률 70% + 수익점수 30%']
    ]
    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text

    doc.add_paragraph()

    doc.add_heading('6.2 수익점수 정규화', level=2)

    p = doc.add_paragraph()
    p.add_run('gain_score = (predicted_gain - (-20)) / (50 - (-20))')
    p.add_run('\n= (예상수익률 + 20) / 70')

    doc.add_paragraph('• 예상수익률 50% 이상 → 수익점수 1.0', style='List Bullet')
    doc.add_paragraph('• 예상수익률 0% → 수익점수 0.286', style='List Bullet')
    doc.add_paragraph('• 예상수익률 -20% 이하 → 수익점수 0.0', style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('6.3 AI 점수 해석', level=2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ['AI 점수', '해석']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    data = [
        ['0.7 이상', '강력 매수 신호 (높은 성공확률 + 높은 예상수익)'],
        ['0.5 ~ 0.7', '관심 종목 (양호한 조건)'],
        ['0.3 ~ 0.5', '중립 (추가 분석 필요)'],
        ['0.3 미만', '약세 신호 (낮은 성공확률)']
    ]
    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text

    doc.add_paragraph()

    # ========================================
    # 7. 시스템 구현
    # ========================================
    doc.add_heading('7. 시스템 구현', level=1)

    doc.add_heading('7.1 기술 스택', level=2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ['구분', '기술']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    data = [
        ['데이터 수집', 'FinanceDataReader, FnGuide 크롤링'],
        ['모델 학습', 'scikit-learn, XGBoost'],
        ['백엔드', 'Python (병렬처리: ThreadPoolExecutor)'],
        ['프론트엔드', 'HTML/CSS/JavaScript (싱글 페이지 앱)']
    ]
    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text

    doc.add_paragraph()

    doc.add_heading('7.2 웹사이트 적용', level=2)

    doc.add_paragraph('• 박스권 돌파 (단순) 스크리너에 AI 점수, 성공확률, 예상수익 컬럼 추가', style='List Bullet')
    doc.add_paragraph('• 실시간 예측: 스크리너 실행 시 각 종목별 ML 모델 추론', style='List Bullet')
    doc.add_paragraph('• 돌파후등락: 돌파일 종가 기준 현재 수익률 표시', style='List Bullet')

    doc.add_paragraph()

    # ========================================
    # 8. 한계점 및 향후 개선
    # ========================================
    doc.add_heading('8. 한계점 및 향후 개선', level=1)

    doc.add_heading('8.1 현재 모델의 한계', level=2)

    doc.add_paragraph('• 회귀 R² 6.87%: 주가의 랜덤성으로 인해 설명력 제한적', style='List Bullet')
    doc.add_paragraph('• Recall 23%: 성공 종목의 일부만 포착 (보수적 접근)', style='List Bullet')
    doc.add_paragraph('• 시장 상황 미반영: 상승장/하락장에 따른 성능 차이 존재', style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('8.2 향후 개선 방향', level=2)

    doc.add_paragraph('• 시장 레짐 변수 추가 (VIX, 투자자별 수급 등)', style='List Bullet')
    doc.add_paragraph('• 섹터별 모델 분리 학습', style='List Bullet')
    doc.add_paragraph('• 딥러닝 모델 (LSTM, Transformer) 실험', style='List Bullet')
    doc.add_paragraph('• 실시간 백테스팅 시스템 구축', style='List Bullet')

    doc.add_paragraph()

    # ========================================
    # 9. 결론
    # ========================================
    doc.add_heading('9. 결론', level=1)

    p = doc.add_paragraph()
    p.add_run('본 프로젝트는 박스권 돌파 패턴에 머신러닝을 적용하여 ')
    p.add_run('투자 의사결정 보조 도구').bold = True
    p.add_run('를 개발하였습니다. ')

    p = doc.add_paragraph()
    p.add_run('주요 성과:')

    doc.add_paragraph('• 14,567개 역사적 돌파 데이터로 학습한 예측 모델 구축', style='List Bullet')
    doc.add_paragraph('• Precision 50.6%: 모델 예측 신뢰도 확보', style='List Bullet')
    doc.add_paragraph('• AI 점수 시스템: 성공확률과 예상수익을 통합한 직관적 지표 제공', style='List Bullet')
    doc.add_paragraph('• 웹 기반 스크리너 연동: 실시간 종목 스크리닝 및 AI 분석', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('\n본 시스템은 투자 참고용이며, 최종 투자 판단은 개인의 책임입니다.').italic = True

    # 저장
    output_path = Path(__file__).parent / "박스권_돌파_ML_모델_보고서_v2.docx"
    doc.save(output_path)
    print(f"보고서 생성 완료: {output_path}")
    return output_path

if __name__ == "__main__":
    create_report()
